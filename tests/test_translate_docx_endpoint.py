from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from translation_service.main import app
from translation_service.models import (
    DocumentPair,
    TranslationUnit,
)

client = TestClient(app)


def create_docx(*paragraphs: str) -> bytes:
    document = Document()

    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    stream = BytesIO()
    document.save(stream)

    return stream.getvalue()


def add_translation(
    db,
    source_text: str,
    target_text: str,
):
    document_pair = DocumentPair(
        source_document="source.docx",
        target_document="target.docx",
    )

    db.add(document_pair)
    db.flush()

    db.add(
        TranslationUnit(
            document_pair_id=document_pair.id,
            source_text=source_text,
            target_text=target_text,
        )
    )

    db.commit()


def test_translate_docx_returns_docx(
    db,
):
    add_translation(
        db,
        "Hei maailma",
        "Hej världen",
    )

    response = client.post(
        "/docx/translate",
        files={
            "file": (
                "source.docx",
                create_docx("Hei maailma"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200

    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    assert "attachment" in response.headers["content-disposition"]


def test_translate_docx_rejects_non_docx_file(
    db,
):
    response = client.post(
        "/docx/translate",
        files={
            "file": (
                "source.txt",
                b"not a docx",
                "text/plain",
            ),
        },
    )

    assert response.status_code == 400

    assert response.json() == {
        "detail": "Only DOCX files are supported",
    }


def test_translate_docx_rejects_invalid_docx(
    db,
):
    response = client.post(
        "/docx/translate",
        files={
            "file": (
                "source.docx",
                b"this is not a valid docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 422

    assert response.json() == {
        "detail": "Invalid DOCX file",
    }


def test_translate_docx_returns_translated_content(
    db,
):
    add_translation(
        db,
        "Hei maailma",
        "Hej världen",
    )

    response = client.post(
        "/docx/translate",
        files={
            "file": (
                "source.docx",
                create_docx(
                    "Hei maailma",
                    "Tuntematon teksti",
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200

    document = Document(BytesIO(response.content))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert paragraphs == [
        "Hej världen",
        "[UNTRANSLATED] Tuntematon teksti",
    ]


def test_translate_docx_preserves_empty_paragraphs(
    db,
):
    add_translation(
        db,
        "Hei maailma",
        "Hej världen",
    )

    add_translation(
        db,
        "Miten voit?",
        "Hur mår du?",
    )

    response = client.post(
        "/docx/translate",
        files={
            "file": (
                "source.docx",
                create_docx(
                    "Hei maailma",
                    "",
                    "Miten voit?",
                ),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200

    document = Document(BytesIO(response.content))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert paragraphs == [
        "Hej världen",
        "",
        "Hur mår du?",
    ]


def test_translate_docx_uses_requested_filename(
    db,
):
    add_translation(
        db,
        "Hei maailma",
        "Hej världen",
    )

    response = client.post(
        "/docx/translate",
        data={
            "output_filename": "my-translation.docx",
        },
        files={
            "file": (
                "source.docx",
                create_docx("Hei maailma"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200

    assert 'filename="my-translation.docx"' in response.headers["content-disposition"]


def test_translate_docx_uses_default_filename(
    db,
):
    add_translation(
        db,
        "Hei maailma",
        "Hej världen",
    )

    response = client.post(
        "/docx/translate",
        files={
            "file": (
                "source.docx",
                create_docx("Hei maailma"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200

    assert 'filename="translated.docx"' in response.headers["content-disposition"]
