from docx import Document

from translation_service.docx_exporter import (
    create_translated_docx,
    translate_paragraphs,
    translate_document,
    ParagraphTranslation,
    build_translated_document,
)
from translation_service.models import DocumentPair, TranslationUnit
from translation_service.translation_status import TranslationStatus


def test_create_translated_docx(tmp_path):
    output_file = tmp_path / "translated.docx"

    create_translated_docx(
        [
            ParagraphTranslation(
                source_text="Hei maailma",
                target_text="Hej världen",
                status=TranslationStatus.TRANSLATED,
            ),
            ParagraphTranslation(
                source_text="Miten voit?",
                target_text="Hur mår du?",
                status=TranslationStatus.TRANSLATED,
            ),
        ],
        output_file,
    )

    document = Document(str(output_file))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert paragraphs == [
        "Hej världen",
        "Hur mår du?",
    ]


def test_translate_paragraphs(db):
    document_pair = DocumentPair(
        source_document="source.docx",
        target_document="target.docx",
    )

    db.add(document_pair)
    db.flush()

    db.add(
        TranslationUnit(
            document_pair_id=document_pair.id,
            source_text="Hei maailma",
            target_text="Hej världen",
        )
    )

    db.commit()

    translated = translate_paragraphs(
        [
            "Hei maailma",
            "Tuntematon teksti",
        ],
        db,
    )

    assert translated[0].target_text == "Hej världen"
    assert translated[0].status == TranslationStatus.TRANSLATED

    assert translated[1].target_text == "Tuntematon teksti"
    assert translated[1].status == TranslationStatus.MISSING


def test_translate_document(
    tmp_path,
    db,
):
    document_pair = DocumentPair(
        source_document="source.docx",
        target_document="target.docx",
    )

    db.add(document_pair)
    db.flush()

    db.add(
        TranslationUnit(
            document_pair_id=document_pair.id,
            source_text="Hei maailma",
            target_text="Hej världen",
        )
    )

    db.commit()

    source_file = tmp_path / "source.docx"
    output_file = tmp_path / "translated.docx"

    source_doc = Document()
    source_doc.add_paragraph("Hei maailma")
    source_doc.add_paragraph("Tuntematon teksti")
    source_doc.save(str(source_file))

    translate_document(
        source_file,
        output_file,
        db,
    )

    translated_doc = Document(str(output_file))

    paragraphs = [paragraph.text for paragraph in translated_doc.paragraphs]

    assert paragraphs == [
        "Hej världen",
        "[UNTRANSLATED] Tuntematon teksti",
    ]


def test_create_translated_docx_marks_missing_translations(
    tmp_path,
):
    output_file = tmp_path / "translated.docx"

    create_translated_docx(
        [
            ParagraphTranslation(
                source_text="Hei maailma",
                target_text="Hej världen",
                status=TranslationStatus.TRANSLATED,
            ),
            ParagraphTranslation(
                source_text="Tuntematon teksti",
                target_text="Tuntematon teksti",
                status=TranslationStatus.MISSING,
            ),
        ],
        output_file,
    )

    document = Document(str(output_file))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert paragraphs == [
        "Hej världen",
        "[UNTRANSLATED] Tuntematon teksti",
    ]


def test_build_translated_document():
    document = build_translated_document(
        [
            ParagraphTranslation(
                source_text="Hei maailma",
                target_text="Hej världen",
                status=TranslationStatus.TRANSLATED,
            ),
            ParagraphTranslation(
                source_text="",
                target_text="",
                status=TranslationStatus.MISSING,
            ),
            ParagraphTranslation(
                source_text="Miten voit?",
                target_text="Hur mår du?",
                status=TranslationStatus.TRANSLATED,
            ),
        ]
    )

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    assert paragraphs == [
        "Hej världen",
        "",
        "Hur mår du?",
    ]
