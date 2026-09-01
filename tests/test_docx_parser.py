from pathlib import Path

from docx import Document

from translation_service.docx_parser import extract_paragraphs


def test_extract_paragraphs(tmp_path: Path) -> None:
    doc_path = tmp_path / "test.docx"

    document = Document()
    document.add_paragraph("Rubrik")
    document.add_paragraph("")
    document.add_paragraph("Första stycket")
    document.add_paragraph("Andra stycket")
    document.save(str(doc_path))
    paragraphs = extract_paragraphs(str(doc_path))

    assert paragraphs == [
        "Rubrik",
        "Första stycket",
        "Andra stycket",
    ]
