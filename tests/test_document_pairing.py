import pytest

from translation_service.document_pairing import (
    pair_paragraphs,
    import_document_pair,
    save_document_pairs,
    import_and_save_document_pair,
)
from docx import Document

from translation_service.models import DocumentPair, TranslationUnit


def test_pair_paragraphs():
    source = [
        "Hei maailma",
        "Miten voit?",
    ]

    target = [
        "Hej världen",
        "Hur mår du?",
    ]

    result = pair_paragraphs(source, target)

    assert result == [
        ("Hei maailma", "Hej världen"),
        ("Miten voit?", "Hur mår du?"),
    ]


def test_rejects_different_number_of_paragraphs():
    source = [
        "Hei maailma",
    ]

    target = [
        "Hej världen",
        "Hur mår du?",
    ]

    with pytest.raises(ValueError):
        pair_paragraphs(source, target)


def test_import_document_pair(tmp_path):
    source_path = tmp_path / "source.docx"
    target_path = tmp_path / "target.docx"

    source_doc = Document()
    source_doc.add_paragraph("Hei maailma")
    source_doc.add_paragraph("Miten voit?")
    source_doc.save(str(source_path))

    target_doc = Document()
    target_doc.add_paragraph("Hej världen")
    target_doc.add_paragraph("Hur mår du?")
    target_doc.save(str(target_path))

    pairs = import_document_pair(
        source_path,
        target_path,
    )

    assert pairs == [
        ("Hei maailma", "Hej världen"),
        ("Miten voit?", "Hur mår du?"),
    ]


def test_save_document_pairs(db):
    count = save_document_pairs(
        "source.docx",
        "target.docx",
        [
            ("Hei maailma", "Hej världen"),
            ("Miten voit?", "Hur mår du?"),
        ],
        db,
    )
    assert count == 2

    document_pairs = db.query(DocumentPair).all()
    assert len(document_pairs) == 1
    assert document_pairs[0].source_document == "source.docx"
    assert document_pairs[0].target_document == "target.docx"

    units = db.query(TranslationUnit).all()
    assert len(units) == 2
    assert units[0].source_text == "Hei maailma"
    assert units[0].target_text == "Hej världen"
    assert units[0].document_pair_id == document_pairs[0].id


def test_import_and_save_document_pair(
    tmp_path,
    db,
):
    source_path = tmp_path / "source.docx"
    target_path = tmp_path / "target.docx"

    source_doc = Document()
    source_doc.add_paragraph("Hei maailma")
    source_doc.add_paragraph("Miten voit?")
    source_doc.save(source_path)

    target_doc = Document()
    target_doc.add_paragraph("Hej världen")
    target_doc.add_paragraph("Hur mår du?")
    target_doc.save(target_path)

    count = import_and_save_document_pair(
        source_path,
        target_path,
        source_path.name,
        target_path.name,
        db,
    )

    assert count == 2

    document_pairs = db.query(DocumentPair).all()
    units = db.query(TranslationUnit).all()

    assert len(document_pairs) == 1
    assert len(units) == 2
