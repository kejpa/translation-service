from pathlib import Path

from docx import Document


def extract_paragraphs(file_path: Path) -> list[str]:
    """
    Extract non-empty paragraphs from a DOCX file
    while preserving the document"""
    document = Document(str(file_path))

    return [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]


def extract_all_paragraphs(
    file_path: Path,
) -> list[str]:
    document = Document(str(file_path))

    return [paragraph.text for paragraph in document.paragraphs]
