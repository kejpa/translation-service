from pathlib import Path

from docx.document import Document as DocxDocument
from docx import Document
from sqlalchemy.orm import Session

from translation_service.docx_parser import extract_all_paragraphs
from translation_service.translation_memory import (
    find_exact_matches,
)
from translation_service.translation_candidates import (
    select_translation_candidate,
)
from dataclasses import dataclass

from translation_service.translation_status import TranslationStatus


@dataclass
class ParagraphTranslation:
    source_text: str
    target_text: str
    status: TranslationStatus


def create_translated_docx(
    paragraphs: list[ParagraphTranslation],
    output_file: Path,
) -> None:
    document = build_translated_document(
        paragraphs,
    )

    document.save(str(output_file))


def translate_paragraphs(
    source_paragraphs: list[str],
    db: Session,
) -> list[ParagraphTranslation]:
    translated_paragraphs: list[ParagraphTranslation] = []

    for paragraph in source_paragraphs:
        if paragraph == "":
            translated_paragraphs.append(
                ParagraphTranslation(
                    source_text="",
                    target_text="",
                    status=TranslationStatus.EMPTY,
                )
            )
            continue
        matches = find_exact_matches(
            paragraph,
            db,
        )

        candidate = select_translation_candidate(matches)

        if candidate is not None:
            translated_paragraphs.append(
                ParagraphTranslation(
                    source_text=paragraph,
                    target_text=candidate.target_text,
                    status=TranslationStatus.TRANSLATED,
                )
            )
        else:
            translated_paragraphs.append(
                ParagraphTranslation(
                    source_text=paragraph,
                    target_text=paragraph,
                    status=TranslationStatus.MISSING,
                )
            )

    return translated_paragraphs


def translate_document(
    source_file: Path,
    output_file: Path,
    db: Session,
) -> None:
    source_paragraphs = extract_all_paragraphs(source_file)

    translations = translate_paragraphs(
        source_paragraphs,
        db,
    )

    create_translated_docx(
        translations,
        output_file,
    )


def build_translated_document(
    paragraphs: list[ParagraphTranslation],
) -> DocxDocument:
    document = Document()

    for paragraph in paragraphs:
        if paragraph.status == TranslationStatus.MISSING:
            document.add_paragraph(f"[UNTRANSLATED] {paragraph.source_text}")
        else:
            document.add_paragraph(paragraph.target_text)

    return document
